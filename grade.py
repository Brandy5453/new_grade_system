
import streamlit as st
import pandas as pd
from difflib import SequenceMatcher
from docx import Document
import re
import io

st.set_page_config(page_title="EUP Tool – Universal Rubric Grader", layout="wide")

st.title("EUP Tool – Universal Rubric Grader")

st.caption("Upload any rubric (Excel/CSV/Word) and a student's DOCX submission to auto-grade against the rubric.")

# ---------- Rubric loaders ----------
def normalize_headers(df: pd.DataFrame) -> pd.DataFrame:
    mapping = {}
    for col in df.columns:
        low = str(col).strip().lower()
        if low in ("criteria", "criterion", "requirement", "item"):
            mapping[col] = "Criteria"
        elif low in ("points", "total points", "score", "max", "max points", "weight"):
            mapping[col] = "Total Points"
        elif "excellent" in low:
            mapping[col] = "Excellent"
        elif "good" in low:
            mapping[col] = "Good"
        elif "poor" in low or "weak" in low:
            mapping[col] = "Poor"
        elif "fair" in low:
            mapping[col] = "Fair"
        elif "description" in low or "descriptor" in low:
            mapping[col] = "Descriptor"
        else:
            mapping[col] = col
    df = df.rename(columns=mapping)
    # If duplicates exist after renaming, keep the first
    df = df.loc[:, ~df.columns.duplicated()]
    if "Criteria" not in df.columns:
        # pick the first object column as Criteria
        for c in df.columns:
            if df[c].dtype == "object":
                df = df.rename(columns={c: "Criteria"})
                break
    if "Total Points" not in df.columns:
        num_cols = [c for c in df.columns if pd.api.types.is_numeric_dtype(df[c])]
        if num_cols:
            df = df.rename(columns={num_cols[0]: "Total Points"})
        else:
            df["Total Points"] = 1.0
    df["Criteria"] = df["Criteria"].astype(str).str.strip()
    df["Total Points"] = pd.to_numeric(df["Total Points"], errors="coerce").fillna(0.0)
    # remove header rows inside data
    df = df[~df["Criteria"].str.lower().str.startswith(("criteria", "criterion"))]
    df = df[df["Criteria"].str.strip() != ""]
    df = df.reset_index(drop=True)
    return df

def load_rubric_excel(file) -> pd.DataFrame:
    # file is a BytesIO
    try:
        xls = pd.ExcelFile(file)
        df = pd.read_excel(xls, xls.sheet_names[0], header=0)
    except Exception:
        # some users upload a single-sheet xlsx created by export; try engine fallback
        df = pd.read_excel(file, header=0)
    return normalize_headers(df)

def load_rubric_csv(file) -> pd.DataFrame:
    df = pd.read_csv(file)
    return normalize_headers(df)

def load_rubric_docx(file) -> pd.DataFrame:
    doc = Document(file)
    tables = doc.tables
    if not tables:
        # fallback parse paragraphs: "<criteria> - <points>"
        rows = []
        for p in doc.paragraphs:
            txt = p.text.strip()
            if not txt:
                continue
            m = re.match(r"(.+?)\s*[-:]\s*([0-9]+(?:\.[0-9]+)?)", txt)
            if m:
                rows.append({"Criteria": m.group(1).strip(), "Total Points": float(m.group(2))})
        df = pd.DataFrame(rows)
        if df.empty:
            raise ValueError("No rubric table/lines found in DOCX.")
        return normalize_headers(df)
    # pick first plausible table
    for t in tables:
        data = []
        cols = len(t.columns)
        header = [cell.text.strip() for cell in t.rows[0].cells]
        looks_header = any("criteria" in c.lower() or "points" in c.lower() for c in header)
        if not looks_header:
            header = [f"Col{i+1}" for i in range(cols)]
        rows_iter = t.rows[1:] if looks_header else t.rows
        for row in rows_iter:
            vals = [cell.text.strip() for cell in row.cells]
            vals = (vals + [""] * cols)[:cols]
            data.append(dict(zip(header, vals)))
        df = pd.DataFrame(data)
        df = normalize_headers(df)
        if not df.empty and "Criteria" in df.columns:
            return df
    raise ValueError("Couldn't parse any table as rubric.")

def load_rubric_any(upload) -> pd.DataFrame:
    name = upload.name.lower()
    if name.endswith((".xlsx", ".xlsm", ".xls")):
        return load_rubric_excel(upload)
    elif name.endswith(".csv"):
        return load_rubric_csv(upload)
    elif name.endswith(".docx"):
        return load_rubric_docx(upload)
    else:
        raise ValueError("Unsupported rubric format. Use Excel/CSV/Word.")

# ---------- Student doc features ----------
def extract_document_features(doc_file):
    doc = Document(doc_file)
    text = "\\n".join(p.text for p in doc.paragraphs)
    headings = [p.text for p in doc.paragraphs if hasattr(p.style, "name") and str(p.style.name).lower().startswith("heading")]
    lists = []
    for p in doc.paragraphs:
        if p.text.strip().startswith(("•", "-", "*", "1.", "a)", "I.")) or \
           (hasattr(p.style, "name") and str(p.style.name).lower().startswith("list")):
            lists.append(p.text)
    formatting = {"bold": [], "italic": [], "underlined": [], "colored_text": []}
    for p in doc.paragraphs:
        for r in p.runs:
            if r.bold: formatting["bold"].append(r.text)
            if r.italic: formatting["italic"].append(r.text)
            if r.underline: formatting["underlined"].append(r.text)
            if hasattr(r.font, "color") and r.font.color and r.font.color.rgb is not None:
                formatting["colored_text"].append(r.text)
    return {
        "text": text,
        "has_images": len(doc.inline_shapes) > 0,
        "headings": headings,
        "lists": lists,
        "formatting": formatting,
        "tables": len(doc.tables),
        "sections": len(doc.sections),
    }

# ---------- Matchers ----------
def similarity_score(a, b):
    return SequenceMatcher(None, a.lower(), b.lower()).ratio()

def keyword_overlap(a, b):
    tok = lambda s: set(re.findall(r"[A-Za-z]{3,}", s.lower()))
    A, B = tok(a), tok(b)
    if not A or not B: return 0.0
    return len(A & B) / len(B)

def match_criteria(student_features, criteria_text):
    if not student_features or not criteria_text:
        return 0.0
    text = student_features["text"]
    if criteria_text.lower() in text.lower():
        return 1.0
    for h in student_features["headings"]:
        if similarity_score(h, criteria_text) > 0.75:
            return 0.95
    ko = keyword_overlap(text, criteria_text)
    sim = similarity_score(text[:5000], criteria_text)
    score = max(0.35 * sim + 0.65 * ko, 0.0)
    return float(min(score, 0.95))

def assess_formatting(features, criteria, max_points):
    c = criteria.lower()
    got = 0.0
    if "bold" in c and features["formatting"]["bold"]: got += 0.4
    if "italic" in c and features["formatting"]["italic"]: got += 0.3
    if "underline" in c and features["formatting"]["underlined"]: got += 0.2
    if "color" in c and features["formatting"]["colored_text"]: got += 0.2
    got = min(got, 1.0)
    return max_points * (got if got > 0 else 0.3)

def assess_images(features, criteria, max_points):
    if not features["has_images"]:
        return 0.0
    if any(x in criteria.lower() for x in ("insert", "add", "include")):
        return max_points
    return max_points * 0.7

def assess_tables(features, criteria, max_points):
    if features["tables"] == 0:
        return 0.0
    if any(x in criteria.lower() for x in ("insert", "add", "include")):
        return max_points
    return max_points * 0.7

def grade_assignment(student_features, rubric_df):
    rows = []
    total_possible = float(rubric_df["Total Points"].sum())
    total_earned = 0.0
    for _, r in rubric_df.iterrows():
        criteria = str(r.get("Criteria", "")).strip()
        max_points = float(r.get("Total Points", 0.0) or 0.0)
        if max_points <= 0 or not criteria:
            rows.append({"Criteria": criteria or "(blank)", "Points Possible": 0.0,
                         "Points Earned": 0.0, "Feedback": "No points allocated", "Match %": "N/A"})
            continue
        c_low = criteria.lower()
        if any(k in c_low for k in ("format", "style", "font", "heading")):
            pts = assess_formatting(student_features, criteria, max_points)
            fb = "Formatting assessed"
        elif any(k in c_low for k in ("image", "picture", "figure", "diagram", "graph")):
            pts = assess_images(student_features, criteria, max_points)
            fb = "Image requirements assessed"
        elif "table" in c_low:
            pts = assess_tables(student_features, criteria, max_points)
            fb = "Table requirements assessed"
        else:
            m = match_criteria(student_features, criteria)
            pts = max_points * m
            fb = f"Content match: {m:.0%}"
        total_earned += pts
        rows.append({
            "Criteria": criteria,
            "Points Possible": round(max_points, 2),
            "Points Earned": round(pts, 2),
            "Feedback": fb,
            "Match %": f"{int(round(100 * (pts / max_points))) if max_points>0 else 0}%"
        })
    return pd.DataFrame(rows), round(total_earned, 2), total_possible

# ---------- UI ----------
with st.expander("1) Upload files", expanded=True):
    rubric_file = st.file_uploader("Rubric (Excel .xlsx / CSV .csv / Word .docx)", type=["xlsx","xlsm","xls","csv","docx"])
    student_doc = st.file_uploader("Student Submission (.docx only)", type=["docx"])

student_name = st.text_input("Student Name")
student_id = st.text_input("Student ID / Index Number")

if rubric_file and student_doc and student_name and student_id:
    try:
        rubric_df = load_rubric_any(rubric_file)
        st.success(f"Rubric loaded with {len(rubric_df)} rows.")
        st.dataframe(rubric_df.head(20))
    except Exception as e:
        st.error(f"Failed to load rubric: {e}")
        st.stop()
    try:
        feats = extract_document_features(student_doc)
    except Exception as e:
        st.error(f"Failed to process student document: {e}")
        st.stop()

    with st.spinner("Grading..."):
        results_df, total_earned, total_possible = grade_assignment(feats, rubric_df)

    st.subheader("Grading Results")
    c1, c2 = st.columns(2)
    with c1:
        st.metric("Total Score", f"{total_earned:.2f}/{total_possible:.2f}")
    with c2:
        pct = (100*total_earned/total_possible) if total_possible else 0
        st.metric("Percentage", f"{pct:.1f}%")

    st.subheader("Detailed Criteria Assessment")
    st.dataframe(results_df)

    # downloads
    csv = results_df.to_csv(index=False).encode("utf-8")
    st.download_button("Download Results (CSV)", csv, file_name=f"grading_{student_id}.csv", mime="text/csv")

    # Simple feedback docx
    from docx import Document as DocxDoc
    fb = DocxDoc()
    fb.add_heading("Automated Grading Report", level=1)
    fb.add_paragraph(f"Student: {student_name} ({student_id})")
    fb.add_paragraph(f"Total: {total_earned}/{total_possible} ({pct:.1f}%)")
    table = fb.add_table(rows=1, cols=4)
    hdr = table.rows[0].cells
    hdr[0].text = "Criteria"
    hdr[1].text = "Points Possible"
    hdr[2].text = "Points Earned"
    hdr[3].text = "Feedback"
    for _, r in results_df.iterrows():
        row = table.add_row().cells
        row[0].text = str(r["Criteria"])
        row[1].text = str(r["Points Possible"])
        row[2].text = str(r["Points Earned"])
        row[3].text = str(r["Feedback"])
    out = io.BytesIO()
    fb.save(out)
    st.download_button("Download Feedback Report (DOCX)", out.getvalue(),
                       file_name=f"grading_report_{student_id}.docx",
                       mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

else:
    st.info("Upload a rubric and a student DOCX, then provide student details to grade.")

st.markdown("---")
st.caption("Tips: Ensure your rubric has at least 'Criteria' and a numeric 'Total Points' column. Other columns (Excellent/Good/Poor) are optional.")
